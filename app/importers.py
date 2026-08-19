"""Parses .docx and .pdf uploads into question records for bulk import.

Expected format (repeated for each question, blank line between questions):

    Q: What is 5 + 3?
    A) 6
    B) 7
    C) 8
    D) 9
    Answer: C
    Explanation: 5 + 3 = 8.

Explanation, Difficulty, and Topic lines are optional. Topic (if given) is
matched against the subject's existing topic names; unmatched or omitted
topics fall back to whatever topic the admin selected on the upload form.
Only text-based PDFs work — scanned/image-only PDFs have no extractable text.
"""
import io
import re

from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader

VALID_DIFFICULTIES = {"easy", "medium", "hard"}

_Q_START = re.compile(r"^\s*Q\d*\s*[:.]\s*(.*)$", re.IGNORECASE)
_CHOICE = re.compile(r"^\s*([A-Da-d])\s*[).]\s*(.+)$")
_ANSWER = re.compile(r"^\s*Answer\s*[:.]\s*([A-Da-d])\b", re.IGNORECASE)
_EXPLANATION = re.compile(r"^\s*Explanation\s*[:.]\s*(.*)$", re.IGNORECASE)
_DIFFICULTY = re.compile(r"^\s*Difficulty\s*[:.]\s*(\w+)", re.IGNORECASE)
_TOPIC = re.compile(r"^\s*Topic\s*[:.]\s*(.+)$", re.IGNORECASE)
_IMAGE_MARKER = re.compile(r"^\x00IMG:(\d+)\x00$")

_IMAGE_EXT_BY_CONTENT_TYPE = {
    "image/png": "png",
    "image/jpeg": "jpg",
    "image/gif": "gif",
    "image/bmp": "bmp",
    "image/tiff": "tiff",
}


class ImportError_(Exception):
    pass


def extract_content(filename, file_bytes):
    """Returns (text, images). images is a list of {"data": bytes, "ext": str},
    referenced from text via \\x00IMG:<index>\\x00 marker lines (docx only —
    PDF images aren't extracted since there's no reliable way to associate
    them with a specific question from text-only page extraction)."""
    lower = filename.lower()
    if lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    if lower.endswith(".pdf"):
        return _extract_pdf(file_bytes), []
    raise ImportError_("Unsupported file type — please upload a .docx or .pdf file.")


def _paragraph_images(doc, paragraph, images):
    """Appends any images embedded in this paragraph to `images` and returns
    marker lines (one per image) to splice into the text stream at this point."""
    markers = []
    for blip in paragraph._p.findall(".//" + qn("a:blip")):
        rId = blip.get(qn("r:embed"))
        if not rId:
            continue
        try:
            part = doc.part.related_parts[rId]
        except KeyError:
            continue
        ext = _IMAGE_EXT_BY_CONTENT_TYPE.get(part.content_type)
        if not ext:
            continue  # skip vector/metafile formats browsers can't render (EMF/WMF)
        images.append({"data": part.blob, "ext": ext})
        markers.append(f"\x00IMG:{len(images) - 1}\x00")
    return markers


def _extract_docx(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ImportError_(f"Could not read this .docx file: {exc}") from exc

    lines = []
    images = []
    for paragraph in doc.paragraphs:
        lines.extend(_paragraph_images(doc, paragraph, images))
        lines.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(lines), images


def _extract_pdf(file_bytes):
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ImportError_(f"Could not read this .pdf file: {exc}") from exc
    pages_text = []
    for page in reader.pages:
        pages_text.append(page.extract_text() or "")
    text = "\n".join(pages_text)
    if not text.strip():
        raise ImportError_(
            "No text could be extracted from this PDF — it may be a scanned image "
            "rather than text. Try a text-based PDF or a .docx file instead."
        )
    return text


def _relocate_dangling_image_markers(lines):
    """An image is commonly placed as its own paragraph directly above the
    question it illustrates, which — since blocks are split purely on 'Q:'
    lines — would otherwise get attributed to the end of the *previous*
    block. Move any marker line that's only followed by blank lines and then
    a 'Q:' line down to just after that 'Q:' line, so it lands inside the
    question it actually belongs to."""
    out = []
    pending = []
    for line in lines:
        if _IMAGE_MARKER.match(line.strip()):
            pending.append(line)
            continue
        if line.strip() == "":
            out.append(line)
            continue
        if _Q_START.match(line) and pending:
            out.append(line)
            out.extend(pending)
            pending = []
            continue
        out.extend(pending)
        pending = []
        out.append(line)
    out.extend(pending)
    return out


def parse_questions_text(raw_text, images=None):
    """Returns (questions, errors). questions: list of dicts ready for insertion
    (minus subject_id/topic_id, which the caller fills in), each with an
    "image_index" key (index into `images`, or None). errors: list of
    human-readable strings describing any block that failed to parse."""
    images = images or []
    lines = raw_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    lines = _relocate_dangling_image_markers(lines)

    blocks = []
    current = None
    for line in lines:
        m = _Q_START.match(line)
        if m:
            if current is not None:
                blocks.append(current)
            current = {"q_first_line": m.group(1), "lines": []}
        elif current is not None:
            current["lines"].append(line)
    if current is not None:
        blocks.append(current)

    questions = []
    errors = []

    for idx, block in enumerate(blocks, start=1):
        question_lines = [block["q_first_line"]] if block["q_first_line"].strip() else []
        choices = {}
        answer = None
        explanation_lines = []
        difficulty = "medium"
        topic_override = None
        image_index = None
        mode = "question"  # question -> choices -> explanation

        for line in block["lines"]:
            stripped = line.strip()
            if not stripped:
                continue

            m_image = _IMAGE_MARKER.match(stripped)
            if m_image:
                candidate = int(m_image.group(1))
                if image_index is None and candidate < len(images):
                    image_index = candidate
                continue  # extra images beyond the first found for this question are ignored

            m_choice = _CHOICE.match(line)
            m_answer = _ANSWER.match(line)
            m_expl = _EXPLANATION.match(line)
            m_diff = _DIFFICULTY.match(line)
            m_topic = _TOPIC.match(line)

            if m_choice:
                letter = m_choice.group(1).upper()
                choices[letter] = m_choice.group(2).strip()
                mode = "choices"
            elif m_answer:
                answer = m_answer.group(1).upper()
                mode = "explanation"
            elif m_diff:
                candidate = m_diff.group(1).strip().lower()
                if candidate in VALID_DIFFICULTIES:
                    difficulty = candidate
            elif m_topic:
                topic_override = m_topic.group(1).strip()
            elif m_expl:
                explanation_lines.append(m_expl.group(1).strip())
            elif mode == "question":
                question_lines.append(stripped)
            elif mode == "explanation":
                explanation_lines.append(stripped)
            # lines appearing after choices but before an Answer: line are ignored
            # (most commonly stray blank/formatting artifacts)

        question_text = " ".join(question_lines).strip()
        explanation = " ".join(explanation_lines).strip()
        preview = question_text[:70] or f"(question #{idx}, no text found)"

        problems = []
        if not question_text:
            problems.append("missing question text")
        missing_letters = [l for l in "ABCD" if l not in choices or not choices[l].strip()]
        if missing_letters:
            problems.append(f"missing choice(s): {', '.join(missing_letters)}")
        if answer is None:
            problems.append("missing 'Answer:' line")
        elif answer not in "ABCD":
            problems.append(f"invalid answer letter '{answer}'")

        if problems:
            errors.append(f"Question #{idx} (\"{preview}\") — {'; '.join(problems)}")
            continue

        questions.append(
            {
                "question_text": question_text,
                "choice_a": choices["A"],
                "choice_b": choices["B"],
                "choice_c": choices["C"],
                "choice_d": choices["D"],
                "correct_choice": answer,
                "explanation": explanation,
                "difficulty": difficulty,
                "topic_override": topic_override,
                "image_index": image_index,
            }
        )

    if not blocks:
        errors.append(
            "No questions were found. Make sure each question starts with a line beginning "
            "'Q:' (or 'Q1:', 'Q2:', ...)."
        )

    return questions, errors
